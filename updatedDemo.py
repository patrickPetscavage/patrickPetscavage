import os
import io
import json
import uuid
import base64
import mimetypes
import time
import re
import threading
from contextlib import nullcontext
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Callable, Optional
from datetime import datetime, timezone

# Extensions treated as images for vision indexing and chat attachments
IMAGE_EXTENSIONS = frozenset(
    {"png", "jpg", "jpeg", "gif", "webp", "bmp", "tif", "tiff"}
)

import streamlit as st
from dotenv import load_dotenv
from openai import (
    OpenAI,
    APIConnectionError,
    APITimeoutError,
    RateLimitError,
    APIStatusError,
)

import chromadb
import pdfplumber
from docx import Document
import openpyxl
try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None

# Load environment variables
load_dotenv()
APP_DIR = os.path.dirname(os.path.abspath(__file__))

# Cached OpenAI client
@st.cache_resource
def get_client(api_key: str):
    return OpenAI(api_key=api_key)

# Persistent Chroma DB client
@st.cache_resource
def get_chroma():
    return chromadb.PersistentClient(path=os.path.join(APP_DIR, ".chroma_db"))

chroma_client = get_chroma()

CHAT_HISTORY_DIR = ".chat_history"
# Stored under CHAT_HISTORY_DIR/images/<scope>/<thread_id>/<file> — JSON holds POSIX-style keys
CHAT_IMAGE_STORE_PREFIX = "images"
INDEXED_DOCS_DIR = ".indexed_docs"

# Use absolute app-rooted paths so Streamlit cwd changes do not break file IO on Windows.
CHAT_HISTORY_DIR = os.path.join(APP_DIR, CHAT_HISTORY_DIR)
INDEXED_DOCS_DIR = os.path.join(APP_DIR, INDEXED_DOCS_DIR)

# Session cost tracking (USD estimates from reported token usage; update if OpenAI changes rates)
SESSION_COST_STATE_KEY = "api_cost_totals_v1"
# https://openai.com/api/pricing/ — values in USD per 1M tokens
PRICE_GPT4O_MINI_INPUT_PER_M = 0.15
PRICE_GPT4O_MINI_OUTPUT_PER_M = 0.60
PRICE_TEXT_EMBEDDING_3_SMALL_PER_M = 0.02
RAG_ANSWER_MAX_TOKENS = 450
RAG_CONTINUATION_MAX_TOKENS = 220
RAG_CONTINUATION_MAX_ROUNDS = 1
MAX_CONTEXT_CHARS = 14000
OPENAI_RETRY_ATTEMPTS = 4
OPENAI_RETRY_BASE_DELAY_S = 1.25
EMBED_BATCH_SIZE = 64
VISION_INDEX_MAX_TOKENS = 420
VISION_MEMORY_MAX_TOKENS = 180
FIGURE_VISION_MAX_TOKENS = 320

# Bulk PDF indexing: per-page vision only when heuristics fire; keep spend low.
PDF_INDEX_VISION_DETAIL = "low"
PDF_INDEX_RENDER_SCALE_SCAN = 2.25
PDF_INDEX_RENDER_SCALE_GRAPHIC = 1.75
PDF_INDEX_SCAN_VISION_MAX_TOKENS = 300
PDF_INDEX_FIGURE_VISION_MAX_TOKENS = 260
# Skip vision for tiny embedded rasters on very text-heavy pages (logos), unless hints apply.
PDF_INDEX_TINY_RASTER_MAX_AREA_RATIO = 0.012
PDF_INDEX_TINY_RASTER_MIN_WORDS = 420
# When a page has multiple embedded rasters, describe each crop (full-page vision misses small images).
PDF_INDEX_MAX_EMBEDDED_IMAGES_PER_PAGE = 12
PDF_INDEX_EMBEDDED_IMAGE_MIN_AREA_RATIO = 0.0035
PDF_INDEX_CROP_RESOLUTION = 110
# Large PDF UX + throughput
LARGE_PDF_PAGE_WARN = 120
LARGE_PDF_PAGE_STRONG_WARN = 400
PDF_INDEX_TEXT_WORKERS = 6
PDF_INDEX_VISION_CONCURRENCY = 3
# Use parallel text extraction when at least this many pages
PDF_INDEX_PARALLEL_TEXT_MIN_PAGES = 14

_GRAPHIC_HINT_RE = re.compile(
    r"\b(fig\.?|figure|diagram|chart|graph|plot|legend|schematic|elevation|"
    r"cross-?section|floor plan|typical|detail|key drivers|schedule|"
    r"isometric|section|plan|detail\s*\d|drawing)\b",
    re.I,
)

# Triggers on-demand source-PDF page rendering + vision during QA (not indexing).
_PAGE_IMAGERY_QUERY_RE = re.compile(
    r"\b(images?|photos?|pictures?|graphics?|illustrations?|maps?|diagrams?|charts?|figures?|"
    r"logos?|vehicles?|equipment|machinery|trucks?|cranes?|construction|"
    r"screenshots?|thumbnails?|below|above|shown|showing|depicts?|depicting|"
    r"looks like|visual(?:ly)?)\b",
    re.I,
)


def fresh_cost_totals() -> dict:
    return {
        "completion": {
            "rag_answer": {"in": 0, "out": 0},
            "vision_index": {"in": 0, "out": 0},
            "vision_memory": {"in": 0, "out": 0},
        },
        "embedding": {"index": 0, "retrieve": 0},
    }


def session_cost_totals() -> dict:
    if SESSION_COST_STATE_KEY not in st.session_state:
        st.session_state[SESSION_COST_STATE_KEY] = fresh_cost_totals()
    return st.session_state[SESSION_COST_STATE_KEY]


def merge_completion_usage(cost_totals: dict, bucket: str, response) -> None:
    u = getattr(response, "usage", None)
    if not u:
        return
    b = cost_totals["completion"].setdefault(bucket, {"in": 0, "out": 0})
    b["in"] += int(getattr(u, "prompt_tokens", 0) or 0)
    b["out"] += int(getattr(u, "completion_tokens", 0) or 0)


def merge_embedding_usage(cost_totals: dict, kind: str, response) -> None:
    u = getattr(response, "usage", None)
    if not u:
        return
    n = int(
        getattr(u, "total_tokens", None)
        or getattr(u, "prompt_tokens", None)
        or 0
    )
    cost_totals["embedding"][kind] = cost_totals["embedding"].get(kind, 0) + n


def completion_bucket_usd(inp: int, out: int) -> float:
    return (
        inp / 1_000_000.0 * PRICE_GPT4O_MINI_INPUT_PER_M
        + out / 1_000_000.0 * PRICE_GPT4O_MINI_OUTPUT_PER_M
    )


def total_session_usd_est(cost_totals: dict) -> float:
    usd = 0.0
    for b in cost_totals["completion"].values():
        usd += completion_bucket_usd(b.get("in", 0), b.get("out", 0))
    emb = cost_totals["embedding"]
    usd += (
        (emb.get("index", 0) + emb.get("retrieve", 0))
        / 1_000_000.0
        * PRICE_TEXT_EMBEDDING_3_SMALL_PER_M
    )
    return usd


def cost_breakdown_select_options(cost_totals: dict) -> list[str]:
    total = total_session_usd_est(cost_totals)
    lines: list[str] = [
        f"Total (session) — ${total:.4f} USD",
    ]
    c = cost_totals["completion"]
    labels = {
        "rag_answer": "Chat answers (gpt-4o-mini, RAG)",
        "vision_index": "Indexing: describe images/diagrams (gpt-4o-mini)",
        "vision_memory": "Chat: image memory summary (gpt-4o-mini)",
    }
    for key, label in labels.items():
        b = c.get(key) or {"in": 0, "out": 0}
        inp, out = b.get("in", 0), b.get("out", 0)
        if inp == 0 and out == 0:
            continue
        u = completion_bucket_usd(inp, out)
        lines.append(f"{label} — ${u:.4f} ({inp:,} in / {out:,} out tok)")
    ei, er = cost_totals["embedding"].get("index", 0), cost_totals["embedding"].get(
        "retrieve", 0
    )
    if ei:
        lines.append(
            f"Embeddings: indexing (text-embedding-3-small) — ${ei / 1_000_000.0 * PRICE_TEXT_EMBEDDING_3_SMALL_PER_M:.4f} ({ei:,} tok)"
        )
    if er:
        lines.append(
            f"Embeddings: search per question (text-embedding-3-small) — ${er / 1_000_000.0 * PRICE_TEXT_EMBEDDING_3_SMALL_PER_M:.4f} ({er:,} tok)"
        )
    if len(lines) == 1:
        lines.append("No API usage recorded yet this session.")
    return lines


def chat_image_key_to_abs(store_key: str) -> str:
    if not store_key:
        return ""
    rel = store_key.replace("/", os.sep)
    return os.path.normpath(os.path.join(CHAT_HISTORY_DIR, rel))


def save_chat_image_to_store(
    scope_key: str,
    thread_id: str,
    original_name: str,
    image_bytes: bytes,
    mime: str,
) -> str:
    """Persist image bytes; return store key (e.g. images/scope/thread/uuid.jpg)."""
    ext = os.path.splitext(original_name)[1].lower()
    if ext not in (
        ".jpg",
        ".jpeg",
        ".png",
        ".gif",
        ".webp",
        ".bmp",
        ".tif",
        ".tiff",
    ):
        m = (mime or "").lower()
        if "png" in m:
            ext = ".png"
        elif "gif" in m:
            ext = ".gif"
        elif "webp" in m:
            ext = ".webp"
        else:
            ext = ".jpg"
    fname = f"{uuid.uuid4().hex}{ext}"
    store_key = f"{CHAT_IMAGE_STORE_PREFIX}/{scope_key}/{thread_id}/{fname}".replace(
        "\\", "/"
    )
    abs_path = chat_image_key_to_abs(store_key)
    os.makedirs(os.path.dirname(abs_path), exist_ok=True)
    with open(abs_path, "wb") as f:
        f.write(image_bytes)
    return store_key


def load_chat_image_bytes(store_key: str) -> Optional[bytes]:
    if not store_key:
        return None
    path = chat_image_key_to_abs(store_key)
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "rb") as f:
            return f.read()
    except OSError:
        return None


def last_user_message_image(
    messages: list,
) -> Optional[tuple[bytes, str]]:
    """Most recent user turn that has a persisted image (for optional re-send to the model)."""
    for m in reversed(messages):
        if m.get("role") != "user":
            continue
        key = m.get("image_store_key")
        mime = m.get("image_mime") or "image/jpeg"
        if not key:
            continue
        data = load_chat_image_bytes(key)
        if data:
            return (data, mime)
    return None


def chat_scope_key(company: str, project: str) -> str:
    return f"{normalize(company)}__{normalize(project)}"

def chat_history_path(company: str, project: str) -> str:
    os.makedirs(CHAT_HISTORY_DIR, exist_ok=True)
    return os.path.join(CHAT_HISTORY_DIR, f"{chat_scope_key(company, project)}.json")

def load_chat_scope_from_disk(company: str, project: str) -> Optional[dict]:
    path = chat_history_path(company, project)
    if not os.path.isfile(path):
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return None

def save_chat_scope_to_disk(company: str, project: str, state: dict) -> None:
    path = chat_history_path(company, project)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(state, f, indent=2, ensure_ascii=False)
    except OSError:
        pass

def ensure_chat_scope_state(company: str, project: str) -> dict:
    sk = chat_scope_key(company, project)
    if "chats_by_scope" not in st.session_state:
        st.session_state.chats_by_scope = {}
    if sk not in st.session_state.chats_by_scope:
        loaded = load_chat_scope_from_disk(company, project)
        if loaded and isinstance(loaded.get("threads"), dict) and loaded["threads"]:
            st.session_state.chats_by_scope[sk] = loaded
            aid = loaded.get("active")
            if aid not in loaded["threads"]:
                st.session_state.chats_by_scope[sk]["active"] = next(iter(loaded["threads"]))
        else:
            tid = str(uuid.uuid4())
            st.session_state.chats_by_scope[sk] = {
                "threads": {tid: {"title": "New chat", "messages": []}},
                "active": tid,
            }
    return st.session_state.chats_by_scope[sk]


def split_user_message_for_display(text: str) -> tuple[str, Optional[str]]:
    """Strip internal [attached image: filename] suffix for UI; return (body, filename)."""
    marker = "\n[attached image:"
    if marker not in text:
        return (text.strip(), None)
    head, tail = text.split(marker, 1)
    fname = tail.strip()
    if fname.endswith("]"):
        fname = fname[:-1].strip()
    return (head.strip() or "(attachment)", fname)


def thread_title_from_message(text: str, max_len: int = 48) -> str:
    body, _fname = split_user_message_for_display(text)
    t = " ".join(body.split())
    if not t:
        return "New chat"
    return (t[: max_len - 1] + "…") if len(t) > max_len else t

# UI reset helper (clears inputs)
def reset_ui():
    keys_to_clear = [
        "chat_input",
        "search_input",
        "temp_input",
    ]

    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

# Normalize company/project names
def normalize(name: str):
    return name.strip().lower().replace(" ", "_")


def normalize_doc_token(name: str) -> str:
    base = os.path.basename(name or "").strip().lower()
    return re.sub(r"[^a-z0-9]+", " ", base).strip()


def now_iso_utc() -> str:
    return datetime.now(timezone.utc).isoformat()


def indexed_doc_scope_dir(company: str, project: str) -> str:
    scope = chat_scope_key(company, project)
    path = os.path.join(INDEXED_DOCS_DIR, scope)
    os.makedirs(path, exist_ok=True)
    return path


# Get or create company collection in Chroma
def get_company_collection(company_id: str):
    safe_name = f"company_{normalize(company_id)}"
    return chroma_client.get_or_create_collection(name=safe_name)

# File readers (PDF, DOCX, Excel, TXT)
def _extract_pdf_page_text(page) -> str:
    """Best-effort machine-readable text extraction for one PDF page."""
    txt = (page.extract_text(layout=True) or "").strip()
    if txt:
        return txt
    txt = (page.extract_text() or "").strip()
    if txt:
        return txt
    words = page.extract_words() or []
    if words:
        return " ".join(w.get("text", "") for w in words if w.get("text", "")).strip()
    return ""


def _page_has_figure_like_images(page, min_area_ratio: float = 0.03) -> bool:
    """
    Heuristic: detect meaningful embedded raster figures/charts on a page.
    Avoid tiny decorative icons by requiring noticeable area.
    """
    images = getattr(page, "images", None) or []
    if not images:
        return False
    page_w = float(getattr(page, "width", 0) or 0)
    page_h = float(getattr(page, "height", 0) or 0)
    page_area = page_w * page_h
    if page_area <= 0:
        return len(images) > 0
    total_img_area = 0.0
    for img in images:
        try:
            x0 = float(img.get("x0", 0) or 0)
            x1 = float(img.get("x1", 0) or 0)
            top = float(img.get("top", 0) or 0)
            bottom = float(img.get("bottom", 0) or 0)
            w = max(0.0, x1 - x0)
            h = max(0.0, bottom - top)
            total_img_area += w * h
        except Exception:
            continue
    return (total_img_area / page_area) >= min_area_ratio


def _page_embedded_raster_area_ratio(page) -> float:
    images = getattr(page, "images", None) or []
    if not images:
        return 0.0
    page_w = float(getattr(page, "width", 0) or 0)
    page_h = float(getattr(page, "height", 0) or 0)
    page_area = page_w * page_h
    if page_area <= 0:
        return 1.0 if images else 0.0
    total = 0.0
    for img in images:
        try:
            x0 = float(img.get("x0", 0) or 0)
            x1 = float(img.get("x1", 0) or 0)
            top = float(img.get("top", 0) or 0)
            bottom = float(img.get("bottom", 0) or 0)
            w = max(0.0, x1 - x0)
            h = max(0.0, bottom - top)
            total += w * h
        except Exception:
            continue
    return total / page_area


def _page_has_any_raster_image(page) -> bool:
    return bool(getattr(page, "images", None))


def _page_text_graphic_hints(page_text: str) -> bool:
    return bool(_GRAPHIC_HINT_RE.search(page_text or ""))


def _page_vector_graphic_suspect(page, page_text: str) -> bool:
    """Heuristic: likely vector art / diagram with little extractable text."""
    words = page.extract_words() or []
    n_words = len(words)
    nc = len((page_text or "").strip())
    page_w = float(getattr(page, "width", 0) or 0)
    page_h = float(getattr(page, "height", 0) or 0)
    page_area = max(page_w * page_h, 1.0)
    density = nc / (page_area / 1000.0) if page_area else 0.0
    if n_words < 28 and nc < 520:
        return True
    if 0 < nc < 900 and density < 11.0:
        return True
    return False


def _pdf_page_should_run_vision_index(
    page,
    page_text: str,
    has_text: bool,
    enable_rich_pdf_vision: bool,
) -> bool:
    """
    Per-page gate: maximize chance to index graphics (including vector-heavy pages),
    while skipping obvious cheap cases (e.g. tiny logo on a long text page).
    """
    has_large_raster = _page_has_figure_like_images(page)
    area_ratio = _page_embedded_raster_area_ratio(page)
    has_any_raster = _page_has_any_raster_image(page)
    mentions_fig = _page_mentions_figure_label(page_text)
    hints = _page_text_graphic_hints(page_text)
    vector = _page_vector_graphic_suspect(page, page_text)
    n_words = len(page.extract_words() or [])

    if not has_text:
        return True

    tiny_only = (
        has_any_raster
        and (not has_large_raster)
        and area_ratio < PDF_INDEX_TINY_RASTER_MAX_AREA_RATIO
        and n_words >= PDF_INDEX_TINY_RASTER_MIN_WORDS
        and not hints
        and not mentions_fig
    )
    if tiny_only:
        return False

    if has_any_raster:
        return True

    if mentions_fig:
        if n_words > 380 and not vector and not hints and not enable_rich_pdf_vision:
            return False
        return True

    if hints and vector:
        return True

    if enable_rich_pdf_vision and hints and n_words < 260:
        return True

    return False


def _page_mentions_figure_label(page_text: str) -> bool:
    t = (page_text or "").lower()
    return ("figure " in t) or ("fig." in t) or (" fig " in t)


def uploaded_file_bytes(file) -> bytes:
    """Read full uploaded file content reliably across Streamlit reruns."""
    try:
        data = file.getvalue()
        if isinstance(data, (bytes, bytearray)):
            return bytes(data)
    except Exception:
        pass
    try:
        file.seek(0)
    except Exception:
        pass
    return file.read() or b""


def persist_uploaded_sources(files, company: str, project: str) -> dict[str, str]:
    """Persist uploaded source files so figure pages can be revisited at QA time."""
    base = indexed_doc_scope_dir(company, project)
    out: dict[str, str] = {}
    for f in files:
        raw = uploaded_file_bytes(f)
        if not raw:
            continue
        safe_name = os.path.basename(f.name)
        path = os.path.join(base, safe_name)
        try:
            with open(path, "wb") as fh:
                fh.write(raw)
            out[f.name] = path
        except OSError:
            continue
    return out


def _render_pdf_page_png_bytes(
    pdfium_doc, page_index: int, scale: float = 2
) -> Optional[bytes]:
    """Render one PDF page to PNG bytes for vision fallback."""
    try:
        page = pdfium_doc[page_index]
        bitmap = page.render(scale=scale)
        image = bitmap.to_pil()
        out = io.BytesIO()
        image.save(out, format="PNG")
        return out.getvalue()
    except Exception:
        return None


def describe_pdf_page_figures_for_rag(
    client: OpenAI,
    image_bytes: bytes,
    filename: str,
    page_num: int,
    cost_totals: Optional[dict] = None,
    image_detail: str = "high",
    max_out_tokens: Optional[int] = None,
    openai_gate: Optional[threading.Semaphore] = None,
) -> str:
    """Extract figure/chart-specific details from a rendered PDF page image."""
    b64 = base64.standard_b64encode(image_bytes).decode("ascii")
    out_cap = max_out_tokens if max_out_tokens is not None else FIGURE_VISION_MAX_TOKENS
    instruction = (
        "This is a full PDF page image. Focus on non-body-text visual content: figures, charts, tables, diagrams, "
        "photos, legends, callouts, and labels. Summarize only visual elements that add meaning beyond plain paragraph text. "
        "Explicitly capture figure/table identifiers when visible (e.g., 'Figure 1', 'Fig. 2', 'Table 3'). "
        "If a caption is visible, quote it briefly. "
        "For charts/diagrams, map colors to categories/components and list visible percentages/values. "
        "If a legend is present, report each legend color and what it represents. "
        "Color accuracy rule: do not guess exact color names when uncertain. "
        "Use cautious wording such as 'dark blue/navy-like' or 'orange-like' and state uncertainty. "
        "If exact values are unclear, say they are unclear. Keep output concise and factual."
    )
    gate_cm = openai_gate if openai_gate is not None else nullcontext()
    with gate_cm:
        res = call_openai_with_retries(
            lambda: client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": instruction},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{b64}",
                                    "detail": image_detail,
                                },
                            },
                        ],
                    }
                ],
                max_tokens=out_cap,
            ),
            operation_name="Figure vision request",
        )
    if cost_totals is not None:
        merge_completion_usage(cost_totals, "vision_index", res)
    out = (res.choices[0].message.content or "").strip()
    if not out:
        return ""
    return f"[FIGURE VISION: {filename} page {page_num}]\n{out}"


def _safe_pdf_image_crop_rect(page, im: dict) -> Optional[tuple[float, float, float, float]]:
    """Clamp embedded-image bbox to page bounds for pdfplumber crop()."""
    try:
        pw = float(page.width or 0)
        ph = float(page.height or 0)
        x0 = max(0.0, float(im.get("x0", 0) or 0))
        top = max(0.0, float(im.get("top", 0) or 0))
        x1 = min(pw, float(im.get("x1", 0) or 0))
        bottom = min(ph, float(im.get("bottom", 0) or 0))
        if x1 <= x0 + 2 or bottom <= top + 2:
            return None
        return (x0, top, x1, bottom)
    except Exception:
        return None


def _vision_index_embedded_raster_crops(
    client: OpenAI,
    page,
    page_num: int,
    filename: str,
    cost_totals: Optional[dict],
    openai_gate: Optional[threading.Semaphore] = None,
) -> str:
    """
    When pdfplumber reports multiple embedded images on one page, run vision on each crop.
    A single full-page vision call often omits smaller photos/diagrams.
    """
    imgs = getattr(page, "images", None) or []
    if len(imgs) < 2:
        return ""

    pw = float(page.width or 0)
    ph = float(page.height or 0)
    page_area = max(pw * ph, 1.0)

    scored: list[tuple[float, int, dict]] = []
    for i, im in enumerate(imgs):
        try:
            rect = _safe_pdf_image_crop_rect(page, im)
            if not rect:
                continue
            x0, top, x1, bottom = rect
            area = max(0.0, x1 - x0) * max(0.0, bottom - top)
            if area / page_area < PDF_INDEX_EMBEDDED_IMAGE_MIN_AREA_RATIO:
                continue
            scored.append((area, i, im))
        except Exception:
            continue

    scored.sort(key=lambda x: -x[0])
    parts: list[str] = []
    for rank, (_area, _i, im) in enumerate(
        scored[:PDF_INDEX_MAX_EMBEDDED_IMAGES_PER_PAGE], start=1
    ):
        rect = _safe_pdf_image_crop_rect(page, im)
        if not rect:
            continue
        try:
            cropped = page.crop(rect)
            pi = cropped.to_image(resolution=PDF_INDEX_CROP_RESOLUTION)
            bio = io.BytesIO()
            pi.original.save(bio, format="PNG")
            raw = bio.getvalue()
        except Exception:
            continue
        if len(raw) < 400:
            continue
        desc = describe_image_for_rag(
            client,
            raw,
            "image/png",
            f"{filename} p{page_num} img{rank}",
            cost_totals=cost_totals,
            image_detail=PDF_INDEX_VISION_DETAIL,
            max_out_tokens=min(VISION_INDEX_MAX_TOKENS, 280),
            openai_gate=openai_gate,
        ).strip()
        if desc:
            parts.append(f"[PDF embedded image {rank} page {page_num}]\n{desc}")

    return "\n\n".join(parts)


def _pdf_extract_range_worker(data: bytes, lo: int, hi: int) -> list[tuple[int, str, list]]:
    """Extract text + image metadata for pages [lo, hi) in one pdfplumber open."""
    rows: list[tuple[int, str, list]] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for idx in range(lo, hi):
            page = pdf.pages[idx]
            imgs = [dict(im) for im in (page.images or [])]
            rows.append((idx, _extract_pdf_page_text(page), imgs))
    return rows


def _pdf_parallel_page_bundles(
    data: bytes,
    n_pages: int,
    progress_cb: Optional[Callable[[float, str], None]] = None,
) -> list[tuple[int, str, list]]:
    """
    Parallel text/image-meta extraction for large PDFs (CPU-bound, separate opens per chunk).
    Returns sorted list of (0-based page index, text, image dict copies).
    """
    if n_pages < PDF_INDEX_PARALLEL_TEXT_MIN_PAGES:
        return _pdf_extract_range_worker(data, 0, n_pages)

    num_workers = min(PDF_INDEX_TEXT_WORKERS, max(2, (n_pages + 79) // 80))
    chunk_size = max(1, (n_pages + num_workers - 1) // num_workers)
    ranges: list[tuple[int, int]] = []
    lo = 0
    while lo < n_pages:
        hi = min(n_pages, lo + chunk_size)
        ranges.append((lo, hi))
        lo = hi

    merged: list[tuple[int, str, list]] = []
    completed = 0
    total = len(ranges)
    with ThreadPoolExecutor(max_workers=min(num_workers, total)) as pool:
        futures = {
            pool.submit(_pdf_extract_range_worker, data, a, b): (a, b)
            for a, b in ranges
        }
        for fut in as_completed(futures):
            merged.extend(fut.result())
            completed += 1
            if progress_cb:
                progress_cb(
                    0.05 + 0.35 * (completed / max(total, 1)),
                    f"Extracting text… block {completed}/{total}",
                )
    merged.sort(key=lambda r: r[0])
    return merged


def read_pdf(
    data: bytes,
    filename: str = "uploaded.pdf",
    client: Optional[OpenAI] = None,
    cost_totals: Optional[dict] = None,
    enable_rich_pdf_vision: bool = False,
    skip_pdf_vision: bool = False,
    progress_cb: Optional[Callable[[float, str], None]] = None,
    openai_vision_gate: Optional[threading.Semaphore] = None,
):
    page_blocks: list[str] = []

    with pdfplumber.open(io.BytesIO(data)) as pdf0:
        n_pages = len(pdf0.pages)

    if n_pages == 0:
        if progress_cb:
            progress_cb(1.0, f"No pages in `{filename}`.")
        return ""

    if progress_cb:
        progress_cb(0.02, f"Reading `{filename}` ({n_pages} pages)…")

    bundles = _pdf_parallel_page_bundles(data, n_pages, progress_cb=progress_cb)
    if progress_cb:
        progress_cb(0.36, f"Text ready: {filename}")

    if skip_pdf_vision or client is None:
        for idx in range(n_pages):
            _, page_text, _imgs = bundles[idx]
            body = (page_text or "").strip() or "(No machine-readable text on this page.)"
            if skip_pdf_vision:
                body = (
                    f"{body}\n"
                    f"(Quick index: PDF vision skipped for speed/cost — "
                    f"re-index with Standard or Full to describe figures/photos.)"
                )
            elif client is None:
                body = (
                    f"{body}\n"
                    "(PDF vision skipped: no OpenAI client configured.)"
                )
            page_blocks.append(f"[PAGE {idx + 1}]\n{body}")
        if progress_cb:
            progress_cb(0.45, f"Text extracted: {filename}")
        return "\n\n".join(page_blocks)

    pdfium_doc = None
    if pdfium is not None:
        try:
            pdfium_doc = pdfium.PdfDocument(data)
        except Exception:
            pdfium_doc = None

    vision_gate = openai_vision_gate

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            bidx = idx - 1
            page_text = bundles[bidx][1] if bidx < len(bundles) else _extract_pdf_page_text(
                page
            )
            has_text = bool(page_text and page_text.strip())

            vis_summary = ""
            should_run_vision = _pdf_page_should_run_vision_index(
                page, page_text, has_text, enable_rich_pdf_vision
            )
            if pdfium_doc is not None and should_run_vision:
                embedded = getattr(page, "images", None) or []
                if len(embedded) >= 2:
                    vis_summary = _vision_index_embedded_raster_crops(
                        client,
                        page,
                        idx,
                        filename,
                        cost_totals,
                        openai_gate=vision_gate,
                    )
                if not vis_summary:
                    scale = (
                        PDF_INDEX_RENDER_SCALE_SCAN
                        if not has_text
                        else PDF_INDEX_RENDER_SCALE_GRAPHIC
                    )
                    rendered = _render_pdf_page_png_bytes(
                        pdfium_doc, idx - 1, scale=scale
                    )
                    if rendered:
                        if not has_text:
                            vis_summary = describe_image_for_rag(
                                client,
                                rendered,
                                "image/png",
                                f"{filename} page {idx}",
                                cost_totals=cost_totals,
                                image_detail=PDF_INDEX_VISION_DETAIL,
                                max_out_tokens=PDF_INDEX_SCAN_VISION_MAX_TOKENS,
                                openai_gate=vision_gate,
                            )
                        else:
                            vis_summary = describe_pdf_page_figures_for_rag(
                                client,
                                rendered,
                                filename,
                                idx,
                                cost_totals=cost_totals,
                                image_detail=PDF_INDEX_VISION_DETAIL,
                                max_out_tokens=PDF_INDEX_FIGURE_VISION_MAX_TOKENS,
                                openai_gate=vision_gate,
                            )

            if not has_text:
                if vis_summary:
                    page_text = (
                        "[VISION FALLBACK: machine-readable text missing]\n"
                        f"{vis_summary}"
                    )
                else:
                    page_text = (
                        "(No machine-readable text found on this page. "
                        "Upload a clearer/scanned page or provide API key for vision fallback.)"
                    )
            elif vis_summary:
                page_text = f"{page_text}\n\n{vis_summary}"
            page_blocks.append(f"[PAGE {idx}]\n{page_text}")

            if progress_cb:
                progress_cb(
                    0.4 + 0.55 * (idx / max(n_pages, 1)),
                    f"Vision pass: page {idx}/{n_pages}",
                )

    if pdfium_doc is not None:
        try:
            pdfium_doc.close()
        except Exception:
            pass
    if progress_cb:
        progress_cb(0.98, f"Finished PDF: {filename}")
    return "\n\n".join(page_blocks)


def get_source_path_for_doc(collection, project: str, source_name: str) -> Optional[str]:
    where_filter = {
        "$and": [
            {"project": project},
            {"source_name": source_name},
        ]
    }
    data = collection.get(where=where_filter, include=["metadatas"])
    metas = data.get("metadatas") or []
    latest_ts = ""
    best_path = None
    for m in metas:
        if not m:
            continue
        p = m.get("source_path")
        if not p:
            continue
        ts = str(m.get("indexed_at") or "")
        if (not best_path) or ts > latest_ts:
            latest_ts = ts
            best_path = p
    return best_path


def extract_figure_context_from_pdf_path(
    pdf_path: str,
    figure_num: str,
    client: OpenAI,
    cost_totals: Optional[dict] = None,
    max_pages: int = 2,
) -> str:
    """
    On-demand PDF vision pass for a specific figure label.
    Returns additional context extracted from rendered page images.
    """
    if not pdfium or (not os.path.isfile(pdf_path)):
        return ""
    try:
        with open(pdf_path, "rb") as fh:
            data = fh.read()
    except OSError:
        return ""
    try:
        pdfium_doc = pdfium.PdfDocument(data)
    except Exception:
        return ""

    label_patterns = [
        f"figure {figure_num}",
        f"fig. {figure_num}",
        f"fig {figure_num}",
    ]

    candidate_pages: list[int] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                t = _extract_pdf_page_text(page).lower()
                if any(p in t for p in label_patterns):
                    candidate_pages.append(i)
    except Exception:
        pass

    if not candidate_pages:
        candidate_pages = [1]

    # Include nearby pages; figures are sometimes separated from text references.
    expanded: list[int] = []
    for p in candidate_pages:
        expanded.extend([p - 1, p, p + 1])
    expanded = sorted({p for p in expanded if p >= 1})[:max_pages]

    chunks: list[str] = []
    for page_num in expanded:
        rendered = _render_pdf_page_png_bytes(pdfium_doc, page_num - 1, scale=2.25)
        if not rendered:
            continue
        try:
            vis = describe_pdf_page_figures_for_rag(
                client,
                rendered,
                os.path.basename(pdf_path),
                page_num,
                cost_totals=cost_totals,
                image_detail="high",
                max_out_tokens=FIGURE_VISION_MAX_TOKENS,
            )
        except Exception:
            vis = ""
        if vis:
            chunks.append(vis)
    try:
        pdfium_doc.close()
    except Exception:
        pass
    return "\n\n".join(chunks)


def extract_figure_images_from_pdf_path(
    pdf_path: str,
    figure_num: str,
    max_pages: int = 2,
) -> list[tuple[bytes, str]]:
    """
    Return rendered PNG images for pages most likely containing Figure N.
    Used as direct image evidence in the final answer call.
    """
    if not pdfium or (not os.path.isfile(pdf_path)):
        return []
    try:
        with open(pdf_path, "rb") as fh:
            data = fh.read()
    except OSError:
        return []
    try:
        pdfium_doc = pdfium.PdfDocument(data)
    except Exception:
        return []

    label_patterns = [
        f"figure {figure_num}",
        f"fig. {figure_num}",
        f"fig {figure_num}",
    ]
    candidate_pages: list[int] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                t = _extract_pdf_page_text(page).lower()
                if any(p in t for p in label_patterns):
                    candidate_pages.append(i)
    except Exception:
        pass
    if not candidate_pages:
        candidate_pages = [1]

    expanded: list[int] = []
    for p in candidate_pages:
        expanded.extend([p - 1, p, p + 1])
    expanded = sorted({p for p in expanded if p >= 1})[:max_pages]

    out: list[tuple[bytes, str]] = []
    for page_num in expanded:
        rendered = _render_pdf_page_png_bytes(pdfium_doc, page_num - 1, scale=3)
        if rendered:
            out.append((rendered, "image/png"))
    try:
        pdfium_doc.close()
    except Exception:
        pass
    return out


def extract_explicit_pages_qa_from_pdf_path(
    pdf_path: str,
    page_nums: list[int],
    client: OpenAI,
    cost_totals: Optional[dict] = None,
    max_pages: int = 3,
    render_scale: float = 3.0,
) -> tuple[str, list[tuple[bytes, str]]]:
    """
    Render specific PDF pages and run figure-style vision for QA (maps, photos, vehicles, etc.).
    Returns (text context block, PNG tuples for the final multimodal answer call).
    """
    if not pdfium or (not os.path.isfile(pdf_path)) or not page_nums:
        return "", []
    try:
        with open(pdf_path, "rb") as fh:
            data = fh.read()
    except OSError:
        return "", []
    pdfium_doc = None
    try:
        pdfium_doc = pdfium.PdfDocument(data)
    except Exception:
        return "", []

    n_doc = len(pdfium_doc)
    pages = [p for p in page_nums[:max_pages] if 1 <= p <= n_doc]
    if not pages:
        try:
            pdfium_doc.close()
        except Exception:
            pass
        return "", []

    text_parts: list[str] = []
    images_out: list[tuple[bytes, str]] = []
    base = os.path.basename(pdf_path)
    try:
        for page_num in pages:
            rendered = _render_pdf_page_png_bytes(
                pdfium_doc, page_num - 1, scale=render_scale
            )
            if not rendered:
                continue
            images_out.append((rendered, "image/png"))
            try:
                vis = describe_pdf_page_figures_for_rag(
                    client,
                    rendered,
                    base,
                    page_num,
                    cost_totals=cost_totals,
                    image_detail="high",
                    max_out_tokens=FIGURE_VISION_MAX_TOKENS,
                )
            except Exception:
                vis = ""
            if vis:
                text_parts.append(vis)
    finally:
        if pdfium_doc is not None:
            try:
                pdfium_doc.close()
            except Exception:
                pass

    header = (
        f"[SOURCE PDF PAGE VISION: {base} — page(s) {', '.join(str(p) for p in pages)}]\n"
        "Describe visible content from these excerpts and the attached page images; "
        "do not claim images are missing if page renders are attached."
    )
    body = "\n\n".join(text_parts)
    if body:
        return f"{header}\n{body}", images_out
    if images_out:
        return header, images_out
    return "", []


def read_docx(data: bytes):
    doc = Document(io.BytesIO(data))
    return "\n".join([p.text for p in doc.paragraphs])

def read_excel(data: bytes):
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    text = ""

    for sheet in wb.worksheets:
        text += f"\n## {sheet.title}\n"
        for row in sheet.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None]
            if vals:
                text += " | ".join(vals) + "\n"

    return text


def guess_image_mime(filename: str, uploaded_type: Optional[str]) -> str:
    if uploaded_type and uploaded_type.startswith("image/"):
        return uploaded_type
    guessed, _ = mimetypes.guess_type(filename)
    if guessed and guessed.startswith("image/"):
        return guessed
    return "image/png"


def describe_image_for_rag(
    client: OpenAI,
    image_bytes: bytes,
    mime: str,
    filename: str,
    cost_totals: Optional[dict] = None,
    image_detail: str = "high",
    max_out_tokens: Optional[int] = None,
    openai_gate: Optional[threading.Semaphore] = None,
) -> str:
    """Turn a diagram or photo into searchable text for embedding (vision)."""
    b64 = base64.standard_b64encode(image_bytes).decode("ascii")
    out_cap = max_out_tokens if max_out_tokens is not None else VISION_INDEX_MAX_TOKENS
    instruction = (
        "Describe this image in detail for document search indexing. "
        "If it is a diagram, floor plan, construction drawing, flowchart, chart, or schematic, "
        "explain components, labels, symbols, relationships, and any visible text, dimensions, or numbers. "
        "Capture explicit identifiers/captions when visible (for example: Figure 1, Fig. 2, Table 3). "
        "For charts/graphs, explicitly map legend colors to categories and include percentages/values for each section when visible. "
        "Color accuracy rule: do not force a specific color name when uncertain; use approximate labels "
        "like 'dark blue/navy-like' and explicitly mark uncertainty. "
        "If a value is not clearly readable, say it is unclear instead of guessing. "
        "If it is a photograph of a site or equipment, note what is shown. Be factual and thorough."
    )
    gate_cm = openai_gate if openai_gate is not None else nullcontext()
    with gate_cm:
        res = call_openai_with_retries(
            lambda: client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": instruction},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime};base64,{b64}",
                                    "detail": image_detail,
                                },
                            },
                        ],
                    }
                ],
                max_tokens=out_cap,
            ),
            operation_name="Vision indexing request",
        )
    if cost_totals is not None:
        merge_completion_usage(cost_totals, "vision_index", res)
    return (res.choices[0].message.content or "").strip()


def brief_image_memory(
    client: OpenAI,
    image_bytes: bytes,
    mime: str,
    cost_totals: Optional[dict] = None,
) -> str:
    """Short vision summary stored with chat history so follow-up questions still have context."""
    b64 = base64.standard_b64encode(image_bytes).decode("ascii")
    instruction = (
        "In 5–8 sentences, summarize what this image shows for chat memory: "
        "type (photo, chart, diagram, plan…), main subjects, readable labels/titles, "
        "key numbers/relationships, and for charts include color-to-category mappings plus visible percentages/values. "
        "No preamble."
    )
    res = call_openai_with_retries(
        lambda: client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": instruction},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"},
                        },
                    ],
                }
            ],
            max_tokens=VISION_MEMORY_MAX_TOKENS,
        ),
        operation_name="Vision memory request",
    )
    if cost_totals is not None:
        merge_completion_usage(cost_totals, "vision_memory", res)
    return (res.choices[0].message.content or "").strip()


def estimate_uploaded_pdf_total_pages(files) -> int:
    """Best-effort page count for large-PDF warnings (opens each PDF briefly)."""
    total = 0
    for f in files:
        if not f.name.lower().endswith(".pdf"):
            continue
        raw = uploaded_file_bytes(f)
        if not raw:
            continue
        try:
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                total += len(pdf.pages)
        except Exception:
            continue
    return total


# Read multiple uploaded files (uses vision for images / diagrams when client is provided)
def read_files(
    files,
    client: Optional[OpenAI] = None,
    cost_totals: Optional[dict] = None,
    enable_rich_pdf_vision: bool = False,
    skip_pdf_vision: bool = False,
    skip_standalone_image_vision: bool = False,
    progress_cb: Optional[Callable[[float, str], None]] = None,
    openai_vision_gate: Optional[threading.Semaphore] = None,
):
    docs: list[dict] = []
    n_files = len(files) or 1

    for fi, f in enumerate(files):

        def _scoped_progress(p: float, msg: str) -> None:
            if not progress_cb:
                return
            base = fi / n_files
            span = 1.0 / n_files
            progress_cb(base + span * min(1.0, max(0.0, p)), msg)

        ext = f.name.split(".")[-1].lower()
        raw = uploaded_file_bytes(f)
        if not raw:
            docs.append(
                {
                    "name": f.name,
                    "text": "(File appears empty or could not be read.)",
                    "status": "empty",
                }
            )
            continue

        try:
            if ext in IMAGE_EXTENSIONS:
                if skip_standalone_image_vision:
                    docs.append(
                        {
                            "name": f.name,
                            "text": (
                                f"[IMAGE / DIAGRAM: {f.name}]\n"
                                "(Quick index: standalone image vision skipped — "
                                "re-index with Standard or Full to describe this image.)"
                            ),
                            "status": "ok",
                        }
                    )
                    if progress_cb:
                        _scoped_progress(1.0, f"{f.name}: skipped (quick mode)")
                    continue
                if not client:
                    docs.append(
                        {
                            "name": f.name,
                            "text": "(Skipped: OpenAI client required to describe images.)",
                            "status": "skipped",
                        }
                    )
                    continue
                mime = guess_image_mime(f.name, getattr(f, "type", None))
                if progress_cb:
                    _scoped_progress(0.1, f"{f.name}: vision…")
                desc = describe_image_for_rag(
                    client,
                    raw,
                    mime,
                    f.name,
                    cost_totals=cost_totals,
                    image_detail=PDF_INDEX_VISION_DETAIL,
                    max_out_tokens=VISION_INDEX_MAX_TOKENS,
                    openai_gate=openai_vision_gate,
                )
                docs.append(
                    {
                        "name": f.name,
                        "text": f"[IMAGE / DIAGRAM: {f.name}]\n{desc}",
                        "status": "ok",
                    }
                )
                if progress_cb:
                    _scoped_progress(1.0, f"{f.name}: done")
            elif ext == "pdf":
                txt = read_pdf(
                    raw,
                    filename=f.name,
                    client=client,
                    cost_totals=cost_totals,
                    enable_rich_pdf_vision=enable_rich_pdf_vision,
                    skip_pdf_vision=skip_pdf_vision,
                    progress_cb=_scoped_progress,
                    openai_vision_gate=openai_vision_gate,
                )
                docs.append(
                    {"name": f.name, "text": f"[DOCUMENT: {f.name}]\n{txt}", "status": "ok"}
                )
            elif ext == "docx":
                if progress_cb:
                    _scoped_progress(0.2, f"{f.name}: reading DOCX…")
                txt = read_docx(raw)
                docs.append(
                    {"name": f.name, "text": f"[DOCUMENT: {f.name}]\n{txt}", "status": "ok"}
                )
                if progress_cb:
                    _scoped_progress(1.0, f"{f.name}: done")
            elif ext in ["xlsx", "xlsm"]:
                if progress_cb:
                    _scoped_progress(0.2, f"{f.name}: reading spreadsheet…")
                txt = read_excel(raw)
                docs.append(
                    {"name": f.name, "text": f"[DOCUMENT: {f.name}]\n{txt}", "status": "ok"}
                )
                if progress_cb:
                    _scoped_progress(1.0, f"{f.name}: done")
            else:
                if progress_cb:
                    _scoped_progress(0.2, f"{f.name}: reading text…")
                txt = raw.decode("utf-8", errors="ignore")
                docs.append(
                    {"name": f.name, "text": f"[DOCUMENT: {f.name}]\n{txt}", "status": "ok"}
                )
                if progress_cb:
                    _scoped_progress(1.0, f"{f.name}: done")
        except Exception as e:
            docs.append(
                {
                    "name": f.name,
                    "text": f"(Indexing error for this file: {e})",
                    "status": "error",
                }
            )

    return docs

# Split text into chunks for embedding
def split_text(text, size=400):
    sections = text.split("\n--- PAGE ")
    chunks = []

    for section in sections:
        words = section.split()
        for i in range(0, len(words), size):
            chunk = " ".join(words[i:i+size])
            chunks.append(chunk)

    return chunks


def docs_to_chunks(
    docs: list[dict],
    chunk_size: int = 400,
) -> tuple[list[str], list[dict]]:
    chunks: list[str] = []
    metadatas: list[dict] = []
    indexed_at = now_iso_utc()
    for d in docs:
        if (d.get("status") or "") != "ok":
            continue
        source_name = d.get("name") or "unknown"
        source_path = d.get("source_path") or ""
        text = (d.get("text") or "").strip()
        if not text:
            continue
        for idx, c in enumerate(split_text(text, size=chunk_size), start=1):
            if not c.strip():
                continue
            chunks.append(c)
            metadatas.append(
                {
                    "source_name": source_name,
                    "source_token": normalize_doc_token(source_name),
                    "source_path": source_path,
                    "chunk_in_source": idx,
                    "indexed_at": indexed_at,
                }
            )
    return chunks, metadatas


def call_openai_with_retries(fn, operation_name: str):
    """Retry transient OpenAI failures with exponential backoff."""
    delay = OPENAI_RETRY_BASE_DELAY_S
    for attempt in range(1, OPENAI_RETRY_ATTEMPTS + 1):
        try:
            return fn()
        except (APIConnectionError, APITimeoutError, RateLimitError) as e:
            if attempt >= OPENAI_RETRY_ATTEMPTS:
                raise RuntimeError(
                    f"{operation_name} failed after {OPENAI_RETRY_ATTEMPTS} attempts: {e}"
                ) from e
            time.sleep(delay)
            delay *= 2
        except APIStatusError as e:
            code = int(getattr(e, "status_code", 0) or 0)
            retryable = code == 429 or code >= 500
            if (not retryable) or attempt >= OPENAI_RETRY_ATTEMPTS:
                raise RuntimeError(
                    f"{operation_name} failed (status {code or 'unknown'}): {e}"
                ) from e
            time.sleep(delay)
            delay *= 2


# Create embeddings using OpenAI
def embed(
    client,
    texts,
    cost_totals: Optional[dict] = None,
    embed_kind: str = "retrieve",
):
    if not texts:
        return []
    vectors = []
    for i in range(0, len(texts), EMBED_BATCH_SIZE):
        batch = texts[i : i + EMBED_BATCH_SIZE]
        res = call_openai_with_retries(
            lambda: client.embeddings.create(
                model="text-embedding-3-small",
                input=batch,
            ),
            operation_name="Embedding request",
        )
        if cost_totals is not None:
            merge_embedding_usage(cost_totals, embed_kind, res)
        vectors.extend([r.embedding for r in res.data])
    return vectors

# Store chunks + embeddings in Chroma
def store_chunks(
    collection,
    project,
    chunks,
    client,
    cost_totals: Optional[dict] = None,
    metadatas: Optional[list[dict]] = None,
):
    embeddings = embed(client, chunks, cost_totals=cost_totals, embed_kind="index")
    ids = [str(uuid.uuid4()) for _ in chunks]
    if metadatas is None:
        metadatas = [{} for _ in chunks]
    full_metadatas = []
    for m in metadatas:
        mm = dict(m or {})
        mm["project"] = project
        full_metadatas.append(mm)

    collection.add(
        ids=ids,
        documents=chunks,
        embeddings=embeddings,
        metadatas=full_metadatas,
    )


def list_indexed_documents(collection, project: str) -> list[dict]:
    data = collection.get(where={"project": project}, include=["metadatas"])
    metas = data.get("metadatas") or []
    by_name: dict[str, dict] = {}
    for m in metas:
        if not m:
            continue
        name = (m.get("source_name") or "(legacy chunks without source)").strip()
        rec = by_name.setdefault(name, {"name": name, "chunks": 0, "last_indexed": ""})
        rec["chunks"] += 1
        ts = str(m.get("indexed_at") or "")
        if ts and (not rec["last_indexed"] or ts > rec["last_indexed"]):
            rec["last_indexed"] = ts
    out = list(by_name.values())
    out.sort(key=lambda r: r["name"].lower())
    return out


def choose_source_hint(query: str, indexed_docs: list[dict]) -> Optional[str]:
    """Infer a likely source document from the question (short names like 'gsi' can match filenames)."""
    junk = {
        "pdf",
        "doc",
        "png",
        "jpg",
        "the",
        "and",
        "for",
        "are",
        "any",
        "what",
        "how",
        "page",
        "pages",
        "tell",
        "about",
        "this",
        "that",
        "with",
        "from",
        "they",
        "them",
        "there",
        "does",
        "document",
    }
    q_tokens = {
        t for t in normalize_doc_token(query).split() if len(t) >= 2
    } - junk
    if not q_tokens:
        return None
    best_name = None
    best_score = 0
    for d in indexed_docs:
        name = d.get("name") or ""
        if name.startswith("(legacy"):
            continue
        base = os.path.splitext(name)[0].lower().replace("_", "-")
        n_tokens = set(normalize_doc_token(name).split()) - junk
        if not n_tokens:
            continue
        score = len(q_tokens & n_tokens)
        compact = re.sub(r"[^a-z0-9]+", "", base)
        for qt in q_tokens:
            if len(qt) >= 3 and qt in compact:
                score += 2
        if score > best_score:
            best_score = score
            best_name = name
    return best_name if best_score >= 2 else None


def query_requests_source_page_imagery(query: str) -> bool:
    return bool(_PAGE_IMAGERY_QUERY_RE.search(query or ""))


def page_numbers_from_query(query: str, max_pages: int = 4) -> list[int]:
    """Parse 1-based page numbers from natural language (e.g. 'page 2', 'p.3')."""
    q = (query or "").lower()
    found: list[int] = []
    for pattern in (
        r"\b(?:page|pg\.?)\s*#?\s*:?\s*(\d{1,4})\b",
        r"\bp\.?\s*#?\s*(\d{1,4})\b",
        r"\[\s*page\s*(\d{1,4})\s*\]",
    ):
        for m in re.finditer(pattern, q):
            n = int(m.group(1))
            if 1 <= n <= 5000:
                found.append(n)
    out: list[int] = []
    seen: set[int] = set()
    for n in found:
        if n not in seen:
            seen.add(n)
            out.append(n)
        if len(out) >= max_pages:
            break
    return out


def _filename_seed_in_prompt(filename: str, prompt: str) -> bool:
    """True if substantive filename tokens appear in the user prompt."""
    stem = os.path.splitext(os.path.basename(filename))[0].lower()
    parts = [p for p in re.split(r"[^a-z0-9]+", stem) if len(p) >= 3]
    pl = (prompt or "").lower()
    return any(p in pl for p in parts)


def pdf_name_for_explicit_page_vision(
    prompt: str,
    source_hint: Optional[str],
    indexed_docs: list[dict],
) -> Optional[str]:
    """
    Pick which indexed PDF to render for explicit page questions.
    If the prompt names a different file than sidebar focus, prefer the named file
    when filename tokens clearly appear in the prompt.
    """
    focused = source_hint if source_hint and source_hint.lower().endswith(".pdf") else None
    inferred = choose_source_hint(prompt, indexed_docs)
    inf_pdf = inferred if inferred and inferred.lower().endswith(".pdf") else None
    if focused and inf_pdf and inf_pdf != focused and _filename_seed_in_prompt(inf_pdf, prompt):
        return inf_pdf
    if focused:
        return focused
    return inf_pdf


def figure_number_from_query(query: str) -> Optional[str]:
    m = re.search(r"\bfig(?:ure)?\.?\s*(\d+)\b", query.lower())
    return m.group(1) if m else None


def figure_focused_chunks(
    collection,
    project: str,
    figure_num: str,
    source_hint: Optional[str] = None,
    max_hits: int = 6,
) -> list[str]:
    """
    Deterministically pull chunks that mention a requested figure label.
    This augments semantic retrieval for prompts like 'explain Figure 1'.
    """
    if source_hint:
        where_filter = {
            "$and": [
                {"project": project},
                {"source_name": source_hint},
            ]
        }
    else:
        where_filter = {"project": project}

    data = collection.get(where=where_filter, include=["documents"])
    docs = data.get("documents") or []
    if not docs:
        return []

    p_exact = f"figure {figure_num}"
    p_figdot = f"fig. {figure_num}"
    p_fig = f"fig {figure_num}"
    scored: list[tuple[int, str]] = []

    for d in docs:
        text = (d or "").lower()
        score = 0
        if p_exact in text:
            score += 8
        if p_figdot in text:
            score += 7
        if p_fig in text:
            score += 6
        if "legend" in text or "color" in text or "colour" in text:
            score += 2
        if "%" in text or "percent" in text or "percentage" in text:
            score += 2
        if score > 0:
            scored.append((score, d))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [doc for _score, doc in scored[:max_hits]]

# Retrieve relevant chunks from vector DB
def retrieve(
    collection,
    project,
    client,
    query,
    k=8,
    cost_totals: Optional[dict] = None,
    source_hint: Optional[str] = None,
):
    query_embedding = embed(client, [query], cost_totals=cost_totals, embed_kind="retrieve")[0]
    if source_hint:
        where_filter = {
            "$and": [
                {"project": project},
                {"source_name": source_hint},
            ]
        }
    else:
        where_filter = {"project": project}

    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=k,
        where=where_filter,
    )

    return results["documents"][0] if results["documents"] else []


def lexical_retrieve(collection, project, query, k=8, source_hint: Optional[str] = None):
    """
    Local fallback retrieval when OpenAI embeddings are unavailable.
    Scores chunks by simple token overlap.
    """
    where_filter = {"project": project}
    if source_hint:
        where_filter = {
            "$and": [
                {"project": project},
                {"source_name": source_hint},
            ]
        }
    data = collection.get(where=where_filter, include=["documents", "metadatas"])
    docs = data.get("documents") or []
    metas = data.get("metadatas") or []
    q_tokens = {t for t in query.lower().split() if len(t) >= 3}
    if not q_tokens:
        return docs[:k]

    scored = []
    for doc in docs:
        d = (doc or "").lower()
        score = sum(1 for t in q_tokens if t in d)
        if score > 0:
            scored.append((score, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [doc for _score, doc in scored[:k]]

def retrieval_query_for_turn(
    prior_messages: list,
    current_question: str,
    max_prior_user_turns: int = 3,
) -> str:
    """Combine recent user questions so follow-ups embed closer to the right document topics."""
    ql = current_question.lower()
    if "figure" in ql or "fig." in ql or "fig " in ql:
        # Figure questions are often very specific; avoid history drift.
        return (
            f"{current_question.strip()}\n"
            "(Include figure references such as Figure N, Fig. N, captions, "
            "chart labels, and page markers [PAGE N].)"
        )
    user_texts = [m["text"] for m in prior_messages if m.get("role") == "user"]
    tail_users = user_texts[-max_prior_user_turns:] if max_prior_user_turns else []
    parts = tail_users + [current_question]
    query = "\n".join(p.strip() for p in parts if p and str(p).strip())

    figure_tokens = ["figure", "fig.", "fig "]
    if any(tok in ql for tok in figure_tokens):
        query += (
            "\n(Include figure references such as Figure N, Fig. N, captions, "
            "chart labels, and page markers [PAGE N].)"
        )
    return query

def format_conversation_history(prior_messages: list, max_messages: int = 12) -> str:
    """Last N messages (before the current user turn) for pronouns and follow-ups."""
    if not prior_messages:
        return ""
    tail = prior_messages[-max_messages:]
    lines = []
    for m in tail:
        role = m.get("role", "")
        text = (m.get("text") or "").strip()
        if not text:
            continue
        label = "User" if role == "user" else "Assistant"
        mem = (m.get("image_memory") or "").strip()
        if role == "user" and mem:
            text = (
                f"{text}\n"
                f"(Visual summary from that message — image not re-sent: {mem})"
            )
        lines.append(f"{label}: {text}")
    return "\n".join(lines)


def trim_context(context: str, max_chars: int = MAX_CONTEXT_CHARS) -> str:
    """Keep prompt size bounded while preserving the most recent/relevant trailing chunks."""
    if not context or len(context) <= max_chars:
        return context
    blocks = context.split("\n\n")
    kept = []
    total = 0
    for block in reversed(blocks):
        add_len = len(block) + (2 if kept else 0)
        if total + add_len > max_chars:
            break
        kept.append(block)
        total += add_len
    kept.reverse()
    return "\n\n".join(kept)

# Generate final answer using retrieved context + optional conversation memory (+ optional image)
def answer(
    client,
    question,
    context,
    conversation_history: Optional[str] = None,
    image_bytes: Optional[bytes] = None,
    image_mime: Optional[str] = None,
    images: Optional[list[tuple[bytes, str]]] = None,
    cost_totals: Optional[dict] = None,
    scoped_document: Optional[str] = None,
):
    history_block = ""
    if conversation_history and conversation_history.strip():
        history_block = f"""
RECENT CONVERSATION (for pronouns and follow-ups; summaries of past images appear here when saved):

{conversation_history}

"""

    system_text = (
        "You are a helpful construction and document assistant.\n"
        "When this request includes a user image, analyze it: read charts, diagrams, plans, labels, "
        "and visible text. Answer the user's question using what you see; do not refuse because "
        "the content is visual or a chart.\n"
        "For pie/bar/line charts, identify legend color -> category mappings and report visible percentages/values.\n"
        "Use the CONTEXT section for facts from indexed project documents. Use the attached image "
        "for anything visible in the image. Combine both when relevant.\n"
        "If the user asks about a specific figure (e.g., Figure 1), first search the provided CONTEXT "
        "for matching figure labels/captions/page markers before saying information is missing.\n"
        "If CONTEXT includes figure-focused excerpts from indexed documents, answer from those excerpts. "
        "Do not ask the user to upload the figure unless they request pixel-level visual details that are not present in text.\n"
        "When reporting colors from figures/charts, be conservative: if color identity is ambiguous, "
        "state an approximate color (e.g., deep blue/navy-like) and mention uncertainty rather than asserting a wrong exact color.\n"
        "If source PDF page images are attached in this turn, prioritize those visual inputs over older text summaries for color mappings.\n"
        "If there is no image in this request but the user asks about a prior attachment, use any "
        "“Visual summary” lines in the recent conversation, and say they can re-attach the file for more detail.\n"
        "If multiple images are provided, treat them as one context (e.g. new upload plus an earlier diagram)."
    )
    if scoped_document:
        system_text += (
            f"\n\nThe user selected **document focus**: `{scoped_document}`. "
            "Treat CONTEXT as coming from that primary source unless a chunk explicitly names another file. "
            "If CONTEXT is empty or clearly unrelated to the question, say the focused document does not "
            "contain matching indexed passages—do not substitute unrelated project documents from memory alone."
        )

    trimmed_context = trim_context(context)
    user_text = f"""{history_block}CONTEXT (indexed documents; may be empty):
{trimmed_context}

USER QUESTION:
{question}"""

    user_content: list = [{"type": "text", "text": user_text}]
    image_parts: list[tuple[bytes, str]] = []
    if images:
        image_parts.extend(images)
    elif image_bytes and image_mime:
        image_parts.append((image_bytes, image_mime))
    for raw, mime in image_parts:
        b64 = base64.standard_b64encode(raw).decode("ascii")
        user_content.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"},
            }
        )

    base_messages = [
        {"role": "system", "content": system_text},
        {"role": "user", "content": user_content},
    ]
    res = call_openai_with_retries(
        lambda: client.chat.completions.create(
            model="gpt-4o-mini",
            messages=base_messages,
            max_tokens=RAG_ANSWER_MAX_TOKENS,
        ),
        operation_name="Chat answer request",
    )
    if cost_totals is not None:
        merge_completion_usage(cost_totals, "rag_answer", res)
    reply = (res.choices[0].message.content or "").strip()
    finish_reason = getattr(res.choices[0], "finish_reason", None)

    # Continue automatically if the model stops due to output token limit.
    continuation_rounds = 0
    while finish_reason == "length" and continuation_rounds < RAG_CONTINUATION_MAX_ROUNDS:
        continuation_rounds += 1
        cont = call_openai_with_retries(
            lambda: client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    *base_messages,
                    {"role": "assistant", "content": reply},
                    {
                        "role": "user",
                        "content": (
                            "Continue exactly where you left off. "
                            "Do not repeat earlier text and do not restart the answer."
                        ),
                    },
                ],
                max_tokens=RAG_CONTINUATION_MAX_TOKENS,
            ),
            operation_name="Chat continuation request",
        )
        if cost_totals is not None:
            merge_completion_usage(cost_totals, "rag_answer", cont)
        next_part = (cont.choices[0].message.content or "").strip()
        if not next_part:
            break
        reply = f"{reply}\n{next_part}".strip()
        finish_reason = getattr(cont.choices[0], "finish_reason", None)

    return reply

# Main Streamlit app
def main():
    st.title("Akilum Cognative Companion")

    api_key = st.sidebar.text_input("OpenAI API Key", type="password")
    if not api_key:
        st.warning("Enter API key")
        return

    client = get_client(api_key)
    cost_totals = session_cost_totals()

    st.sidebar.header("Session cost (estimate)")
    st.sidebar.metric(
        "Total (USD)",
        f"${total_session_usd_est(cost_totals):.4f}",
        help=(
            "Sum of this browser session’s tracked OpenAI calls, using list prices for "
            "gpt-4o-mini and text-embedding-3-small. Your invoice may differ."
        ),
    )
    st.sidebar.selectbox(
        "Breakdown",
        options=cost_breakdown_select_options(cost_totals),
        key="session_cost_breakdown_select",
    )
    st.sidebar.caption(
        f"Rates used: gpt-4o-mini ${PRICE_GPT4O_MINI_INPUT_PER_M}/${PRICE_GPT4O_MINI_OUTPUT_PER_M} per M in/out; "
        f"text-embedding-3-small ${PRICE_TEXT_EMBEDDING_3_SMALL_PER_M}/M. See openai.com/api/pricing."
    )
    if st.sidebar.button("Reset session cost", key="reset_session_cost_btn"):
        st.session_state[SESSION_COST_STATE_KEY] = fresh_cost_totals()
        st.rerun()

    # (Only showing the UPDATED company section — everything else stays the same)

    # COMPANY MANAGEMENT SECTION
    st.sidebar.header("🏢 Company")

    COMPANY_FILE = ".companies.json"

    def load_companies():
        if not os.path.exists(COMPANY_FILE):
            return []
        with open(COMPANY_FILE, "r") as f:
            return json.load(f)

    def save_companies(companies):
        with open(COMPANY_FILE, "w") as f:
            json.dump(sorted(list(set(companies))), f)

    companies = load_companies()

    # Initialize session state
    if "company" not in st.session_state:
        st.session_state["company"] = None

    if "show_company_input" not in st.session_state:
        st.session_state.show_company_input = False

    col1, col2 = st.sidebar.columns([3, 1])

    with col1:
        selected_company = st.selectbox(
            "Select Company",
            companies if companies else ["No companies yet"],
            key="company_select_clean"
        )

    with col2:
        if st.button("➕"):
            st.session_state.show_company_input = True

    # Inline creation
    if st.session_state.get("show_company_input"):
        new_company = st.sidebar.text_input("New company name")

        if st.sidebar.button("Create"):
            if new_company.strip():
                norm = normalize(new_company)
                companies.append(norm)
                save_companies(companies)

                # IMPORTANT FIX: actually create collection immediately
                get_company_collection(norm)

                st.session_state["company"] = norm
                st.session_state.show_company_input = False
                reset_ui()
                st.rerun()

    # Set active company
    if selected_company != "No companies yet":
        if st.session_state.get("company") != selected_company:
            st.session_state["company"] = selected_company

    company = st.session_state["company"]

    # FIX: safety check so UI doesn't crash
    if company:
        st.sidebar.success(f"Active Company: {company}")
    else:
        st.sidebar.warning("No company selected")

    # FIX: prevent crash when no company exists
    if not company:
        st.stop()

    # FIX: collection must be created AFTER company is guaranteed
    collection = get_company_collection(company)

    # PROJECT MANAGEMENT SECTION

    PROJECT_FILE = ".projects.json"

    def load_projects():
        if not os.path.exists(PROJECT_FILE):
            return {}
        with open(PROJECT_FILE, "r") as f:
            return json.load(f)

    def save_projects(data):
        with open(PROJECT_FILE, "w") as f:
            json.dump(data, f, indent=2)

    # Load all projects (scoped by company)
    all_projects = load_projects()

    if "project" not in st.session_state:
        st.session_state["project"] = None

    if company not in all_projects:
        all_projects[company] = []
        save_projects(all_projects)

    company_projects = all_projects.get(company, [])

    st.sidebar.header("Project")

    # UI state
    if "show_project_input" not in st.session_state:
        st.session_state.show_project_input = False

    col1, col2 = st.sidebar.columns([3, 1])

    with col1:
        selected_project = st.selectbox(
            "Select Project",
            company_projects if company_projects else ["No projects yet"],
            key=f"project_select_{company}"
    )

    with col2:
        if st.button("➕", key=f"add_project_{company}"):
            st.session_state.show_project_input = True

    # Inline project creation
    if st.session_state.get("show_project_input"):
        new_project = st.sidebar.text_input("Create Project", key="new_project_input")

        if st.sidebar.button("Create Project"):
            if new_project.strip():
                company_projects.append(new_project)

                all_projects[company] = sorted(list(set(company_projects)))
                save_projects(all_projects)

                st.session_state["project"] = new_project
                st.session_state.show_project_input = False
                reset_ui()
                st.rerun()

    # Set active project (NO rerun here)
    if selected_project != "No projects yet":
        if st.session_state.get("project") != selected_project:
            st.session_state["project"] = selected_project

    # Initialize default project
    if "project" not in st.session_state or not st.session_state["project"]:
        if company_projects:
            st.session_state["project"] = company_projects[0]

    project = st.session_state["project"]

    # Safety
    if not project:
        st.sidebar.warning("No project selected")
        st.stop()

    st.sidebar.success(f"Active Project: {project}")

    # Keep your existing document UI exactly the same
    indexed_docs = list_indexed_documents(collection, project)

    with st.sidebar.expander("Indexed documents in this project", expanded=False):
        if not indexed_docs:
            st.caption("No indexed chunks yet for this project.")
        else:
            st.caption(f"{len(indexed_docs)} document source(s) indexed.")
            for d in indexed_docs:
                ts = d.get("last_indexed") or "unknown time"
                st.caption(f"- {d['name']} — {d['chunks']} chunks — last: {ts}")

    focus_options = ["(Auto)"] + [
        d["name"] for d in indexed_docs if not d["name"].startswith("(legacy")
    ]

    doc_focus_choice = st.sidebar.selectbox(
        "Document focus for QA",
        options=focus_options,
        key=f"doc_focus_{chat_scope_key(company, project)}",
        help=(
            "Auto tries to infer the document from your question. "
            "Pick a specific PDF when asking about figures/captions—or phrases like “page 2” with "
            "graphics/photos—so retrieval (and optional page rendering) stay on that file."
        ),
    )

    # FILE UPLOAD + INDEXING SECTION
    files = st.sidebar.file_uploader(
        "Upload documents & images/diagrams",
        type=[
            "pdf",
            "docx",
            "xlsx",
            "xlsm",
            "txt",
            "png",
            "jpg",
            "jpeg",
            "gif",
            "webp",
            "bmp",
            "tif",
            "tiff",
        ],
        accept_multiple_files=True,
        key=f"uploader_{company}_{project}",
        help="PDFs/DOCX/Excel/TXT are read as text. Images are described with vision and indexed like text.",
    )
    index_profile = st.sidebar.radio(
        "Indexing mode",
        options=("quick", "standard", "full"),
        format_func=lambda v: {
            "quick": "Quick — text only (fastest, lowest cost)",
            "standard": "Standard — smart PDF vision + images",
            "full": "Full — richer PDF figure coverage (higher cost)",
        }[v],
        index=1,
        help=(
            "Quick extracts PDF text in parallel when large, but skips all vision (PDF figures and "
            "uploaded images). Standard uses gated low-detail vision on graphic-like PDF pages and "
            "describes standalone images. Full relaxes gating so more text-heavy pages get figure vision."
        ),
        key=f"index_profile_{chat_scope_key(company, project)}",
    )
    skip_pdf_vision = index_profile == "quick"
    skip_image_vision = index_profile == "quick"
    enable_rich_pdf_vision = index_profile == "full"

    if st.sidebar.button("Index Documents"):
        if not files:
            st.error("Upload files first")
        else:
            pdf_pages = estimate_uploaded_pdf_total_pages(files)
            if pdf_pages >= LARGE_PDF_PAGE_STRONG_WARN:
                st.warning(
                    f"This batch includes about **{pdf_pages}** PDF pages. "
                    "Indexing may take a long time and many API calls in Standard/Full. "
                    "Consider **Quick** first to get text online, then re-run Standard/Full on a subset."
                )
            elif pdf_pages >= LARGE_PDF_PAGE_WARN:
                st.info(
                    f"Large PDF load (~{pdf_pages} pages). Progress is shown below; "
                    "Streamlit cannot cancel a run mid-flight."
                )

            persisted_map = persist_uploaded_sources(files, company, project)
            prog = st.progress(0.0)
            status = st.empty()
            vision_gate = (
                threading.Semaphore(PDF_INDEX_VISION_CONCURRENCY)
                if client and not (skip_pdf_vision and skip_image_vision)
                else None
            )

            def _index_progress(p: float, msg: str) -> None:
                prog.progress(min(1.0, max(0.0, p)))
                status.caption(msg)

            docs = read_files(
                files,
                client,
                cost_totals,
                enable_rich_pdf_vision=enable_rich_pdf_vision,
                skip_pdf_vision=skip_pdf_vision,
                skip_standalone_image_vision=skip_image_vision,
                progress_cb=_index_progress,
                openai_vision_gate=vision_gate,
            )
            prog.progress(1.0)
            status.empty()
            for d in docs:
                nm = d.get("name")
                if nm in persisted_map:
                    d["source_path"] = persisted_map[nm]
            chunks, chunk_metas = docs_to_chunks(docs)
            if not chunks:
                st.error("No readable content was found in the uploaded files.")
            else:
                try:
                    store_chunks(
                        collection,
                        project,
                        chunks,
                        client,
                        cost_totals,
                        metadatas=chunk_metas,
                    )
                    ok_docs = [d["name"] for d in docs if (d.get("status") or "") == "ok"]
                    st.success(
                        f"Indexed {len(chunks)} chunks from {len(ok_docs)} document(s)."
                    )
                    st.rerun()
                except Exception as e:
                    st.error(
                        "Indexing failed while contacting OpenAI. "
                        "Please retry in a moment and check your network/API key."
                    )
                    st.caption(f"Details: {e}")

    # CHAT THREADS (per company + project): switch, history, disk persistence
    chat_state = ensure_chat_scope_state(company, project)
    threads: dict = chat_state["threads"]
    active_id = chat_state["active"]

    st.sidebar.header("Chats")

    thread_ids = list(threads.keys())
    if active_id not in threads and thread_ids:
        chat_state["active"] = thread_ids[0]
        active_id = chat_state["active"]
        save_chat_scope_to_disk(company, project, chat_state)
        st.rerun()

    def _thread_label(tid: str) -> str:
        t = threads.get(tid, {})
        title = t.get("title") or "Chat"
        n = len(t.get("messages") or [])
        return f"{title} ({n // 2} turns)" if n else title

    col_new, col_del = st.sidebar.columns(2)
    with col_new:
        if st.button("➕ New chat", key=f"new_chat_{chat_scope_key(company, project)}"):
            tid = str(uuid.uuid4())
            threads[tid] = {"title": "New chat", "messages": []}
            chat_state["active"] = tid
            save_chat_scope_to_disk(company, project, chat_state)
            reset_ui()
            st.rerun()
    with col_del:
        if len(thread_ids) > 1 and st.button("🗑️ Delete", key=f"del_chat_{chat_scope_key(company, project)}"):
            del threads[active_id]
            chat_state["active"] = next(iter(threads.keys()))
            save_chat_scope_to_disk(company, project, chat_state)
            reset_ui()
            st.rerun()

    sel_index = thread_ids.index(active_id) if active_id in thread_ids else 0
    picked = st.sidebar.selectbox(
        "Active chat",
        options=thread_ids,
        index=sel_index,
        format_func=_thread_label,
        key=f"chat_pick_{chat_scope_key(company, project)}",
    )
    if picked != active_id:
        chat_state["active"] = picked
        save_chat_scope_to_disk(company, project, chat_state)
        st.rerun()

    active_thread = threads[chat_state["active"]]
    messages = active_thread.setdefault("messages", [])

    if "image_upload_nonce" not in st.session_state:
        st.session_state.image_upload_nonce = 0

    scope_k = chat_scope_key(company, project)
    pin_key = f"pin_chat_image_{scope_k}"

    chat_image = st.sidebar.file_uploader(
        "Attach diagram to next message (optional)",
        type=["png", "jpg", "jpeg", "gif", "webp", "bmp", "tif", "tiff"],
        key=f"chat_attach_{scope_k}_{st.session_state.image_upload_nonce}",
        help=(
            "Images are saved under .chat_history/images/ so they show in the thread after refresh. "
            "Optional: re-send the latest saved image on every message (see checkbox below)."
        ),
    )

    pin_chat_image = st.sidebar.checkbox(
        "Re-send latest saved diagram with every message",
        key=pin_key,
        help=(
            "Includes the most recent image from this chat in each model request until you upload a new one. "
            "Uses vision input tokens on every turn while enabled."
        ),
    )

    for msg in messages:
        role = msg.get("role")
        text = msg.get("text") or ""
        if role == "user":
            body, att_name = split_user_message_for_display(text)
            store_key = msg.get("image_store_key")
            with st.chat_message("user"):
                st.write(body)
                blob = load_chat_image_bytes(store_key) if store_key else None
                if blob:
                    st.image(io.BytesIO(blob), caption=att_name or "Saved attachment")
                elif att_name:
                    st.caption(
                        f"Attached: {att_name} (file not on disk — send again to restore preview)."
                    )
        else:
            st.chat_message(role).write(text)

    # CHAT INPUT + RAG FLOW
    prompt = st.chat_input("Ask something...", key="chat_input")

    if prompt:
        img_bytes: Optional[bytes] = None
        img_mime: Optional[str] = None
        if chat_image is not None:
            img_bytes = chat_image.read()
            img_mime = guess_image_mime(
                chat_image.name, getattr(chat_image, "type", None)
            )

        with st.chat_message("user"):
            st.write(prompt)
            if img_bytes:
                st.image(io.BytesIO(img_bytes), caption=chat_image.name)

        call_images: Optional[list[tuple[bytes, str]]] = None
        if img_bytes and img_mime:
            call_images = [(img_bytes, img_mime)]
        elif pin_chat_image:
            pinned = last_user_message_image(messages)
            if pinned:
                call_images = [pinned]
        figure_page_images: list[tuple[bytes, str]] = []

        rag_query = retrieval_query_for_turn(messages, prompt)
        source_hint = (
            doc_focus_choice
            if doc_focus_choice != "(Auto)"
            else choose_source_hint(rag_query, indexed_docs)
        )
        if img_bytes:
            rag_query = f"{rag_query}\n[user attached an image/diagram]"
        elif call_images:
            rag_query = f"{rag_query}\n[last saved chat image included in model context]"
        scoped_for_answer = (
            source_hint if doc_focus_choice != "(Auto)" else None
        )
        context = []
        fig_num = figure_number_from_query(prompt)
        retrieval_note: Optional[str] = None
        # Figure queries can be served from deterministic figure-matching chunks + on-demand
        # OpenAI vision on source PDF pages, avoiding fragile embedding retrieval.
        if not fig_num:
            try:
                context = retrieve(
                    collection,
                    project,
                    client,
                    rag_query,
                    k=8,
                    cost_totals=cost_totals,
                    source_hint=source_hint,
                )
                if source_hint and not context:
                    context = retrieve(
                        collection,
                        project,
                        client,
                        rag_query,
                        k=24,
                        cost_totals=cost_totals,
                        source_hint=source_hint,
                    )
                if source_hint and not context:
                    context = lexical_retrieve(
                        collection,
                        project,
                        rag_query,
                        k=12,
                        source_hint=source_hint,
                    )
                if source_hint and not context:
                    retrieval_note = (
                        f"No indexed passages matched in focused document `{source_hint}`. "
                        "Answer from attached source PDF page images when present; do not pull in other files."
                    )
                    st.warning(retrieval_note)
                elif not source_hint and not context:
                    context = lexical_retrieve(
                        collection,
                        project,
                        rag_query,
                        k=10,
                        source_hint=None,
                    )
            except Exception as e:
                st.error("OpenAI retrieval failed for this message. Please retry.")
                st.caption(f"Retrieval details: {e}")
                return
        if source_hint:
            st.caption(f"Using document focus: {source_hint}")

        context_text = "\n\n".join(context)
        if retrieval_note:
            context_text = f"[RETRIEVAL NOTE]\n{retrieval_note}\n\n{context_text}"

        page_qa_images: list[tuple[bytes, str]] = []
        explicit_pages = page_numbers_from_query(rag_query)
        if explicit_pages and query_requests_source_page_imagery(prompt) and not fig_num:
            pdf_pick = pdf_name_for_explicit_page_vision(
                prompt, source_hint, indexed_docs
            )
            if pdf_pick:
                if source_hint and pdf_pick != source_hint:
                    st.caption(
                        f"Rendering pages from **{pdf_pick}** (filename tokens in your question)."
                    )
                src_path_pages = get_source_path_for_doc(
                    collection, project, pdf_pick
                )
                if src_path_pages:
                    page_ctx, page_imgs = extract_explicit_pages_qa_from_pdf_path(
                        src_path_pages,
                        explicit_pages,
                        client,
                        cost_totals=cost_totals,
                    )
                    if page_ctx or page_imgs:
                        if page_ctx:
                            context_text = f"{page_ctx}\n\n{context_text}"
                        page_qa_images = page_imgs
                        st.caption(
                            f"Source PDF page vision: **{pdf_pick}** page(s) "
                            f"{', '.join(str(p) for p in explicit_pages)} "
                            f"({len(page_imgs)} render(s))."
                        )

        if fig_num:
            fig_hits = figure_focused_chunks(
                collection,
                project,
                fig_num,
                source_hint=source_hint,
                max_hits=6,
            )
            if fig_hits:
                fig_block = "\n\n".join(fig_hits[:4])
                context_text = (
                    f"[FIGURE-FOCUSED CONTEXT: Figure {fig_num}]\n"
                    f"{fig_block}\n\n"
                    f"{context_text}"
                )
                st.caption(
                    f"Added figure-focused context for Figure {fig_num} ({len(fig_hits)} chunk match(es))."
                )
            if source_hint and source_hint.lower().endswith(".pdf"):
                src_path = get_source_path_for_doc(collection, project, source_hint)
                if src_path:
                    ondemand_vis = extract_figure_context_from_pdf_path(
                        src_path,
                        fig_num,
                        client,
                        cost_totals=cost_totals,
                        max_pages=2,
                    )
                    if ondemand_vis:
                        context_text = (
                            f"[ON-DEMAND FIGURE VISION FROM SOURCE PDF: Figure {fig_num}]\n"
                            f"{ondemand_vis}\n\n"
                            f"{context_text}"
                        )
                        st.caption(
                            f"Added on-demand vision context from source PDF for Figure {fig_num}."
                        )
                    figure_page_images = extract_figure_images_from_pdf_path(
                        src_path,
                        fig_num,
                        max_pages=2,
                    )
                    if figure_page_images:
                        st.caption(
                            f"Attached {len(figure_page_images)} source PDF page image(s) for Figure {fig_num}."
                        )

        pdf_evidence_images = [*figure_page_images, *page_qa_images]
        if pdf_evidence_images:
            if call_images:
                call_images = [*call_images, *pdf_evidence_images]
            else:
                call_images = pdf_evidence_images

        history_text = format_conversation_history(messages)
        try:
            reply = answer(
                client,
                prompt,
                context_text,
                conversation_history=history_text or None,
                images=call_images,
                cost_totals=cost_totals,
                scoped_document=scoped_for_answer,
            )
        except Exception as e:
            st.error("OpenAI response generation failed for this message. Please retry.")
            st.caption(f"Generation details: {e}")
            return

        st.chat_message("assistant").write(reply)

        if active_thread.get("title") == "New chat":
            active_thread["title"] = thread_title_from_message(prompt)

        user_line = prompt
        if img_bytes:
            user_line = f"{prompt}\n[attached image: {chat_image.name}]"
        user_entry: dict = {"role": "user", "text": user_line}
        if img_bytes:
            sk_img = save_chat_image_to_store(
                scope_k,
                chat_state["active"],
                chat_image.name,
                img_bytes,
                img_mime or "image/jpeg",
            )
            user_entry["image_store_key"] = sk_img
            user_entry["image_mime"] = img_mime or "image/jpeg"
            user_entry["image_memory"] = brief_image_memory(
                client, img_bytes, user_entry["image_mime"], cost_totals=cost_totals
            )
        messages.append(user_entry)
        messages.append({"role": "assistant", "text": reply})
        save_chat_scope_to_disk(company, project, chat_state)

        if img_bytes:
            st.session_state.image_upload_nonce += 1
            st.rerun()

#START APP
if __name__ == "__main__":
    main()
